"""
Document parsing and chunking module.

Supports:
- DOCX (Microsoft Word)
- IPYNB (Jupyter Notebooks)
- TXT/MD (Plain text files)

Chunking uses semantic block splitting with configurable overlap.
"""

import json
import re
import logging
from pathlib import Path
from typing import List, Tuple

from docx import Document

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT CLEANING
# ═══════════════════════════════════════════════════════════════════════════════


def _clean(text: str) -> str:
    """
    Normalize whitespace and remove excessive blank lines.
    
    Args:
        text: Raw text input
        
    Returns:
        Cleaned text with normalized whitespace
    """
    # Replace non-breaking spaces
    text = text.replace("\u00a0", " ")
    # Normalize horizontal whitespace (preserve newlines)
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# BLOCK SPLITTING
# ═══════════════════════════════════════════════════════════════════════════════


def _split_blocks(text: str) -> List[str]:
    """
    Split text into semantic blocks (paragraphs).
    
    Splits on blank lines to preserve natural document structure.
    
    Args:
        text: Cleaned text
        
    Returns:
        List of non-empty text blocks
    """
    blocks = re.split(r"\n\s*\n", text)
    return [b.strip() for b in blocks if b.strip()]


# ═══════════════════════════════════════════════════════════════════════════════
# CHUNKING
# ═══════════════════════════════════════════════════════════════════════════════


def _chunk_by_char_budget(
    blocks: List[str],
    max_chars: int,
    overlap_chars: int,
) -> List[str]:
    """
    Combine blocks into chunks respecting character budget.
    
    Strategy:
    1. Accumulate blocks until max_chars would be exceeded
    2. Flush buffer as a chunk
    3. Add overlap from previous chunk to next
    
    Args:
        blocks: List of text blocks
        max_chars: Maximum characters per chunk
        overlap_chars: Characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    chunks: List[str] = []
    buffer: List[str] = []
    current_length = 0

    def flush() -> None:
        nonlocal buffer, current_length
        if not buffer:
            return
        chunk = "\n\n".join(buffer).strip()
        if chunk:
            chunks.append(chunk)
        buffer = []
        current_length = 0

    for block in blocks:
        block_len = len(block)

        # Handle oversized blocks by hard splitting
        if block_len > max_chars:
            flush()  # Flush any pending content
            for i in range(0, block_len, max_chars):
                sub = block[i : i + max_chars].strip()
                if sub:
                    chunks.append(sub)
            continue

        # Check if adding this block would exceed budget
        # +2 accounts for "\n\n" separator
        if current_length + block_len + 2 > max_chars:
            flush()

        buffer.append(block)
        current_length += block_len + 2

    flush()  # Don't forget the last buffer

    # Apply overlap between chunks
    if overlap_chars > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            overlap = prev[-overlap_chars:] if len(prev) > overlap_chars else prev
            overlapped.append((overlap + "\n\n" + chunks[i]).strip())
        chunks = overlapped

    return chunks


# ═══════════════════════════════════════════════════════════════════════════════
# FILE PARSERS
# ═══════════════════════════════════════════════════════════════════════════════


def parse_docx(path: str) -> str:
    """
    Extract text from a Microsoft Word document.
    
    Args:
        path: Path to .docx file
        
    Returns:
        Extracted and cleaned text
    """
    try:
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        return _clean("\n".join(lines))
    except Exception as e:
        logger.error(f"Failed to parse DOCX {path}: {e}")
        raise


def parse_ipynb(path: str) -> str:
    """
    Extract text from a Jupyter Notebook.
    
    Preserves cell type information (markdown vs code).
    
    Args:
        path: Path to .ipynb file
        
    Returns:
        Extracted and formatted text
    """
    try:
        data = json.loads(Path(path).read_text(encoding="utf-8", errors="ignore"))
        output_parts = []

        for cell in data.get("cells", []):
            cell_type = cell.get("cell_type")
            source = "".join(cell.get("source", [])).strip()

            if not source:
                continue

            if cell_type == "markdown":
                output_parts.append(f"## Notebook Markdown\n{source}")
            elif cell_type == "code":
                output_parts.append(f"## Notebook Code\n```python\n{source}\n```")
            else:
                output_parts.append(source)

        return _clean("\n\n".join(output_parts))
    except Exception as e:
        logger.error(f"Failed to parse IPYNB {path}: {e}")
        raise


def parse_text_like(path: str) -> str:
    """
    Read plain text files (txt, md, etc.).
    
    Args:
        path: Path to text file
        
    Returns:
        Cleaned text content
    """
    try:
        return _clean(Path(path).read_text(encoding="utf-8", errors="ignore"))
    except Exception as e:
        logger.error(f"Failed to read text file {path}: {e}")
        raise


# ═══════════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════════


def load_text_for_file(filepath: str) -> Tuple[str, str]:
    """
    Load and parse a file based on its extension.
    
    Args:
        filepath: Path to the file
        
    Returns:
        Tuple of (extracted_text, file_type)
        
    Supported extensions:
        - .docx → "docx"
        - .ipynb → "ipynb"
        - .txt, .md, etc. → "text"
    """
    p = Path(filepath)
    ext = p.suffix.lower()

    if ext == ".docx":
        return parse_docx(filepath), "docx"
    if ext == ".ipynb":
        return parse_ipynb(filepath), "ipynb"
    return parse_text_like(filepath), "text"


def chunk_text(
    text: str,
    max_chars: int = 1400,
    overlap_chars: int = 250,
) -> List[str]:
    """
    Split text into overlapping chunks for embedding.
    
    Args:
        text: Full document text
        max_chars: Maximum characters per chunk (default: 1400)
        overlap_chars: Overlap between chunks (default: 250)
        
    Returns:
        List of text chunks ready for embedding
    """
    if not text or not text.strip():
        return []

    blocks = _split_blocks(text)
    return _chunk_by_char_budget(
        blocks,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
    )

import json
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document


def _clean(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_blocks(text: str) -> List[str]:
    # split by blank lines (better semantic blocks)
    return [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]


def _chunk_by_char_budget(blocks: List[str], max_chars: int, overlap_chars: int) -> List[str]:
    chunks: List[str] = []
    buf: List[str] = []
    cur = 0

    def flush():
        nonlocal buf, cur
        if not buf:
            return
        chunk = "\n\n".join(buf).strip()
        if chunk:
            chunks.append(chunk)
        buf = []
        cur = 0

    for b in blocks:
        if len(b) > max_chars:
            # hard split giant blocks
            for i in range(0, len(b), max_chars):
                sub = b[i : i + max_chars].strip()
                if sub:
                    chunks.append(sub)
            continue

        if cur + len(b) + 2 > max_chars:
            flush()

        buf.append(b)
        cur += len(b) + 2

    flush()

    # simple overlap
    if overlap_chars > 0 and len(chunks) > 1:
        out = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            overlap = prev[-overlap_chars:] if len(prev) > overlap_chars else prev
            out.append((overlap + "\n\n" + chunks[i]).strip())
        chunks = out

    return chunks


def parse_docx(path: str) -> str:
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    return _clean("\n".join(lines))


def parse_ipynb(path: str) -> str:
    data = json.loads(Path(path).read_text(encoding="utf-8", errors="ignore"))
    out = []
    for cell in data.get("cells", []):
        ctype = cell.get("cell_type")
        src = "".join(cell.get("source", [])).strip()
        if not src:
            continue

        if ctype == "markdown":
            out.append(f"## Notebook Markdown\n{src}")
        elif ctype == "code":
            out.append(f"## Notebook Code\n```python\n{src}\n```")
        else:
            out.append(src)
    return _clean("\n\n".join(out))


def parse_text_like(path: str) -> str:
    return _clean(Path(path).read_text(encoding="utf-8", errors="ignore"))


def load_text_for_file(filepath: str) -> Tuple[str, str]:
    """
    Returns (text, kind)
    """
    p = Path(filepath)
    ext = p.suffix.lower()

    if ext == ".docx":
        return parse_docx(filepath), "docx"
    if ext == ".ipynb":
        return parse_ipynb(filepath), "ipynb"
    return parse_text_like(filepath), "text"


def chunk_text(text: str, max_chars: int = 1400, overlap_chars: int = 250) -> List[str]:
    blocks = _split_blocks(text)
    return _chunk_by_char_budget(blocks, max_chars=max_chars, overlap_chars=overlap_chars)
